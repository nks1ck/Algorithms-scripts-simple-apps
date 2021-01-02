import docx
from docx.shared import RGBColor, Pt

# Get message from fake message
fake_text = docx.Document('fakeMessage.docx')
fake_list = []

for paragraph in fake_text.paragraphs:
    fake_list.append(paragraph.text)

# Get text from real message
real_text = docx.Document('realMessage.docx')
real_list = []

for paragraph in real_text.paragraphs:
    if len(paragraph.text) != 0:
        real_list.append(paragraph.text)

# Get template form style, font and etc.
doc = docx.Document('template.docx')


doc.add_heading('Ilsaf Nabiullin', 0)
subtitle = doc.add_heading('Nks1ckk & co', 1)
subtitle.alignment = 1
doc.add_heading('', 1)
doc.add_paragraph('2 Jan 2021')
doc.add_paragraph('')


def set_spacing(paragraph):
    """set space between paragraphs"""
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)


LENGTH_REAL = len(real_list)
COUNT_REAL = 0

for line in fake_list:
    if COUNT_REAL < LENGTH_REAL and line == "":
        paragraph = doc.add_paragraph(real_list[COUNT_REAL])
        paragraph_index = len(doc.paragraphs) - 1

        # Set white color to font
        run = doc.paragraphs[paragraph_index].runs[0]
        font = run.font
        font.color.rgb = RGBColor(255, 255, 255)

        COUNT_REAL += 1
    else:
        paragraph = doc.add_paragraph(line)

    set_spacing(paragraph)

doc.save('new.docx')

print("done!")